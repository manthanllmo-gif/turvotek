import os
import docx
from pypdf import PdfReader

def extract_docx(docx_path, output_md_path):
    print(f"Extracting docx from {docx_path} to {output_md_path}")
    doc = docx.Document(docx_path)
    content = []
    for para in doc.paragraphs:
        content.append(para.text)
    
    # Handle tables in docx
    for table in doc.tables:
        content.append("\n--- TABLE ---")
        for row in table.rows:
            row_data = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            content.append(" | ".join(row_data))
        content.append("---------------\n")
        
    with open(output_md_path, "w", encoding="utf-8") as f:
        f.write("\n".join(content))

def extract_pdf(pdf_path, output_md_path):
    print(f"Extracting pdf from {pdf_path} to {output_md_path}")
    reader = PdfReader(pdf_path)
    content = []
    for i, page in enumerate(reader.pages):
        content.append(f"\n--- PAGE {i+1} ---")
        content.append(page.extract_text() or "")
    
    with open(output_md_path, "w", encoding="utf-8") as f:
        f.write("\n".join(content))

if __name__ == "__main__":
    extract_docx(
        r"d:\turvotek\assetspictures\Website Content Requirements.docx",
        r"d:\turvotek\website_content_requirements.md"
    )
    extract_pdf(
        r"d:\turvotek\assetspictures\Solar Accessories 06.06.2026.pdf",
        r"d:\turvotek\solar_accessories_catalog.md"
    )
    print("Done extracting!")
